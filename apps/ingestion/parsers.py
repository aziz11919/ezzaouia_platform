"""
Parsers pour extraction de données depuis PDF, Word et Excel.
"""
import logging
import pandas as pd
import PyPDF2
from docx import Document

logger = logging.getLogger('apps')


def parse_excel(filepath):
    """
    Lit un fichier Excel et retourne une liste de dicts.
    Détecte automatiquement la feuille de production.
    """
    try:
        xl = pd.ExcelFile(filepath)
        logger.info(f"Feuilles trouvées : {xl.sheet_names}")

        # Prendre la première feuille par défaut
        df = pd.read_excel(filepath, sheet_name=0)

        # Nettoyer : supprimer lignes vides
        df = df.dropna(how='all')

        # Normaliser les noms de colonnes
        df.columns = [str(c).strip() for c in df.columns]

        records = df.to_dict(orient='records')
        logger.info(f"Excel : {len(records)} lignes extraites")
        return records, None

    except Exception as e:
        logger.error(f"Erreur Excel : {e}")
        return [], str(e)


def parse_pdf(filepath):
    """
    Extrait le texte brut d'un PDF.
    Retourne le texte complet pour indexation RAG.
    """
    try:
        text = ""
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {i+1} ---\n{page_text}"

        logger.info(f"PDF : {len(reader.pages)} pages extraites")
        return text, None

    except Exception as e:
        logger.error(f"Erreur PDF : {e}")
        return "", str(e)


def parse_word(filepath):
    """
    Extrait le texte d'un fichier Word (.docx).
    Retourne le texte complet pour indexation RAG.
    """
    try:
        doc = Document(filepath)
        text = ""

        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # Extraire aussi les tableaux
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    text += row_text + "\n"

        logger.info(f"Word : {len(doc.paragraphs)} paragraphes extraits")
        return text, None

    except Exception as e:
        logger.error(f"Erreur Word : {e}")
        return "", str(e)